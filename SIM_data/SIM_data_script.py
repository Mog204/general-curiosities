# -*- coding: utf-8 -*-
"""
Created on Monday 11 November 2025

@author: Imogen S

---- Description ----
This script reads default Simulation files created for Cabrini Hospital, and logs them into a searchable database based on key features.
"""

from pathlib import Path

# Path for sim files, including any grouped folders
path_sim_files = Path("C:/Users/imy1/Documents/Work/Mama/Clean SIM data")

import os
import numpy as np
from docx import Document
import ntpath
import collections
import re
from datetime import datetime

# optional: support .doc via textract (requires installation and native deps)
try:
    import textract
    HAS_TEXTRACT = True
except Exception:
    HAS_TEXTRACT = False


def findWholeWord(w):
    return re.compile(r'\b({0})\b'.format(w), flags=re.IGNORECASE).search


def getMetaData(doc):
    metadata = {}
    prop = doc.core_properties
    metadata["author"] = prop.author
    metadata["category"] = prop.category
    metadata["comments"] = prop.comments
    metadata["content_status"] = prop.content_status
    metadata["created"] = prop.created
    metadata["identifier"] = prop.identifier
    metadata["keywords"] = prop.keywords
    metadata["last_modified_by"] = prop.last_modified_by
    metadata["language"] = prop.language
    metadata["modified"] = prop.modified
    metadata["subject"] = prop.subject
    metadata["title"] = prop.title
    metadata["version"] = prop.version
    return metadata

def get_title(doc_path, fname):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if 'title' in searchtext:
            index = i
            break
    if index != None:
        if all_strips[index].replace("title","").replace(":","").replace("*","").strip() != "":
            title_text = all_strips[index].replace("title","").replace(":","").replace("*","").strip()
        else:
            title_text = all_strips[index+1].replace("title","").replace(":","").replace("*","").strip()
        return title_text, 0
    else:
        return fname, 1
    
def get_date(doc_path, meta_date):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if findWholeWord('date')(searchtext)  and i<50 :
            index = i
            break
    if index != None:
        if all_strips[index].replace("date","").replace(":","").replace("*","").strip() != "":
            search_date = all_strips[index].replace("date","").replace(":","").replace("*","").strip()
        elif all_strips[index+1].replace("date","").replace(":","").replace("*","").strip() != "":
            search_date = all_strips[index+1].replace("date","").replace(":","").replace("*","").strip()
        else:
            search_date = meta_date
        
        # Parse the input string into a datetime object
        # %d for day, %m for month, %y for two-digit year

        if search_date[-4:].isnumeric() == False:
                search_date = search_date[:-2] + "20" + search_date[-2:]

        if '/' not in search_date: #search_date.replace(" ","").isalpha() == True:
            search_date = meta_date.strftime("%d/%m/%Y") if meta_date != None else None
        
        date_object = datetime.strptime(search_date, "%d/%m/%Y")

        # %Y for four-digit year
        search_date_yyyy = date_object.strftime("%d/%m/%Y")

        return search_date_yyyy, 0

    else:
        fixed_meta_date = meta_date.strftime("%d/%m/%Y") if meta_date != None else None
        return fixed_meta_date, 1
    
def get_author(doc_path, meta_author):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if findWholeWord('author')(searchtext)  and i<50 :
            index = i
            break
    if index != None:
        if all_strips[index].replace("author","").replace(":","").replace("*","").strip() != "":
            search_author = all_strips[index].replace("author","").replace("*","").replace(":","").strip()
        elif all_strips[index+1].replace("author","").replace(":","").replace("*","").strip() != "":
            search_author = all_strips[index+1].replace("author","").replace(":","").replace("*","").strip()
        else:
            search_author = meta_author
        return search_author, 0
    else:
        if meta_author == 'python-docx':
            return None, 1
        return meta_author, 1
    
def get_target(doc_path):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if findWholeWord('target')(searchtext) or findWholeWord('audience')(searchtext) or 'players' in searchtext \
            or findWholeWord('developed for')(searchtext) and i<50 :
            index = i
            break
    if index != None:
        if all_strips[index].replace("target","").replace("audience","").replace("intended","").replace("players","").replace(":","").replace("*","").strip() != "":
            search_target = all_strips[index].replace("target","").replace("audience","").replace("intended","").replace("players","").replace(":","").replace("*","").strip()
        elif all_strips[index+1].replace("target","").replace("audience","").replace("intended","").replace("players","").replace(":","").replace("*","").strip() != "":
            search_target = all_strips[index+1].replace("target","").replace("audience","").replace("intended","").replace("players","").replace(":","").replace("*","").strip()
        else:
            search_target = None
        if findWholeWord('role')(search_target) or len(search_target) >100:
            search_target = None
        return search_target, 0
    else:
        return None, 1
    

# Objectives
def get_objectives(doc_path):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if 'objective' in searchtext or 'goal' in searchtext:
            index = i+1
            test_ind = index
            while document.paragraphs[test_ind].text.lower().strip() != "" and test_ind < len(document.paragraphs)-1:
                test_ind +=1
            
            end_index = test_ind
            break
    if index != None:
        extracted_text = [document.paragraphs[i].text.lower().strip() for i in range(index, end_index)]
        search_obj = "\n".join(extracted_text)
        if len(search_obj) > 1000:
            search_obj = None
        return search_obj, 0
    else:
        return None, 1

# Reviewer
def get_reviewer(doc_path):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if findWholeWord('reviewer')(searchtext):
            index = i+1
            break
    if index != None:
        if all_strips[index].replace("reviewer","").replace(":","").replace("*","").strip() != "":
            search_reviewer = all_strips[index].replace("reviewer","").replace(":","").replace("*","").strip()
        elif all_strips[index+1].replace("reviewer","").replace(":","").replace("*","").strip() != "":
            search_reviewer = all_strips[index+1].replace("reviewer","").replace(":","").replace("*","").strip()
        else:
            search_reviewer = None
        return search_reviewer, 0
    else:
        return None, 1
    
# Difficulty
def get_difficulty(doc_path):
    document = Document(doc_path)
    all_strips = []
    index = None
    for i,para in enumerate(document.paragraphs):
        searchtext = para.text.lower().strip()
        all_strips.append(searchtext)
        if 'difficulty' in searchtext:
            index = i
            break
    if index != None:
        if all_strips[index].replace("difficulty","").replace(":","").replace("*","").strip() != "":
            search_difficulty = all_strips[index].replace("difficulty","").replace(":","").replace("*","").strip()
        elif all_strips[index+1].replace("difficulty","").replace(":","").replace("*","").strip() != "":
            search_difficulty = all_strips[index+1].replace("difficulty","").replace(":","").replace("*","").strip()
        else:
            search_difficulty = None

        if '/' not in search_difficulty or len(search_difficulty) >1000:
            search_difficulty = None
        return search_difficulty, 0
    else:
        return None, 1

def extract_titles(root_dir, output_dir=None, save_csv=True):
    sim_titles = []
    types = []
    rows=[]
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for fname in filenames:
            file_title = os.path.splitext(fname)[0]
            ext = os.path.splitext(fname)[1].lower()
            if ext in (".docx", ".doc"):
                try:
                    src_path = os.path.join(dirpath, fname)
                    sim_titles.append(get_title(src_path, file_title)[0])
                    types.append(get_title(src_path, file_title)[1])
                except Exception as e:
                    # record failure and continue
                    rows.append({"source_path": src_path, "status": "error", "message": str(e), "output_path": None})
                    continue
    return sim_titles, types, rows

sim_titles_test, types, rows = extract_titles(path_sim_files)

import fitz  # PyMuPDF
from docx import Document

def extract_pdf_to_word(pdf_path, word_path):
    """
    Extracts text from a PDF file and saves it into a Word document.

    Args:
        pdf_path (str): The path to the input PDF file.
        word_path (str): The path to the output Word document.
    """
    try:
        # Open the PDF document
        pdf_document = fitz.open(pdf_path)

        # Create a new Word document
        word_document = Document()

        # Iterate through each page of the PDF
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")  # Extract text from the page

            # Add the extracted text as a paragraph to the Word document
            word_document.add_paragraph(text)

        # Save the Word document
        word_document.save(word_path)
        print(f"Text successfully extracted from '{pdf_path}' to '{word_path}'")

    except Exception as e:
        print(f"An error occurred: {e}")

def convert_to_docs(root_dir):
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for fname in filenames:
            ext = os.path.splitext(fname)[1].lower()
            if ext == ".pdf":
                try:
                    src_path = os.path.join(dirpath, fname)
                    base_name = os.path.splitext(fname)[0]
                    output_path = os.path.join(dirpath, base_name + ".docx")
                    extract_pdf_to_word(src_path, output_path)
                except Exception as e:
                    print(f"An error occurred while processing '{src_path}': {e}")
convert_to_docs(path_sim_files)

import pandas as pd

features = ["Source", "Title", "Author", "Date", "Attendees", "Reviewer", "Objectives", "Difficulty"]
def extract_info(root_dir, output_dir=None, save_csv=True):
    sim_titles = []
    types = []
    rows=[]
    

    len_docs = 0
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for fname in filenames:
            file_title = os.path.splitext(fname)[0]
            ext = os.path.splitext(fname)[1].lower()
            if ext in (".docx", ".doc"):
                len_docs += 1

    
    info_df = pd.DataFrame(index=range(len_docs),columns = features)
    
    ind = 0
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for fname in filenames:
            file_title = os.path.splitext(fname)[0]
            ext = os.path.splitext(fname)[1].lower()
            if ext in (".docx", ".doc"):
                try:
                    src_path = os.path.join(dirpath, fname)
                    doc = Document(src_path)

                    # Extract source path
                    info_df.at[ind, "Source"] = src_path

                    # Extract title
                    info_df.at[ind,"Title"] = get_title(src_path, file_title)[0]
                    ##sim_titles.append(get_title(src_path, file_title)[0])
                    ##types.append(get_title(src_path, file_title)[1])

                    # Extract metadata
                    metadata_dict = getMetaData(doc)


                    # Extract date
                    info_df.at[ind,"Date"]= get_date(src_path, metadata_dict["created"])[0]

                    # Extract author
                    info_df.at[ind,"Author"]= get_author(src_path, metadata_dict["author"])[0]
                    
                    # Extract Target / attendees
                    info_df.at[ind,"Attendees"]= get_target(src_path)[0]

                    # Extract Objectives
                    info_df.at[ind,"Objectives"]= get_objectives(src_path)[0]

                    # Extract Reviewer
                    info_df.at[ind,"Reviewer"]= get_reviewer(src_path)[0]

                    # Extract Difficulty
                    info_df.at[ind,"Difficulty"]= get_difficulty(src_path)[0]

                    ind += 1
                    ##dates.append(metadata_dict["created"])


                    
                except Exception as e:
                    # record failure and continue
                    rows.append({"source_path": src_path, "status": "error", "message": str(e), "output_path": None})
                    continue
    return info_df

info_df_test = extract_info(path_sim_files)

# save to searchable table in excel

import xlsxwriter
import pandas as pd

# create directory for excel
directory_path = path_sim_files / "Excel summaries"
os.makedirs(directory_path, exist_ok=True)

def get_letter_from_position(position, uppercase=False):
    if uppercase:
        # 'A' is Unicode 65, so add (position - 1) to 65
        return chr(65 + (position - 1))
    else:
        # 'a' is Unicode 97, so add (position - 1) to 97
        return chr(97 + (position - 1))

writer = pd.ExcelWriter(Path(directory_path / f'{datetime.strftime(datetime.today(), '%Y-%m-%d')} SIM table.xlsx'), engine='xlsxwriter')


# 3. Write the DataFrame data to XlsxWriter
info_df_test.to_excel(writer, sheet_name='Sims', index=False)


# 4. Get the xlsxwriter workbook and worksheet objects
workbook = writer.book
worksheet = writer.sheets['Sims']

# Add a format for text wrapping
wrap_format = workbook.add_format({'text_wrap': True})

# 5. Add the Excel table structure
# The range corresponds to the data in the DataFrame (3 columns, 3 rows + header)
end_letter = get_letter_from_position(len(info_df_test.columns), uppercase=True)
end_row = len(info_df_test) + 1  # +1 for header row
table_range = f'A1:{end_letter}{end_row}'

print(table_range)

columns_dict = [{'header': col} for col in info_df_test.columns]

worksheet.add_table(table_range, {'columns': columns_dict})

for col_num, value in enumerate(info_df_test.columns.values):
    worksheet.set_column(col_num, col_num, 30, wrap_format)
# 6. Close the Pandas Excel writer and output the Excel file
writer.close()